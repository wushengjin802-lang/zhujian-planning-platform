from __future__ import annotations

import re
import time
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from sqlalchemy import delete, select, text
from sqlalchemy.orm import Session

from app.db.models import DocumentChunk, ParseJob, ProjectDocument
from app.services.storage import storage


def chunk_text(content: str, size: int = 700) -> list[str]:
    text_value = re.sub(r"\s+", " ", content).strip()
    if not text_value:
        return ["未识别到可解析正文，需人工补充资料摘要。"]
    return [text_value[index : index + size] for index in range(0, len(text_value), size)]


def chunk_sections(sections: list[tuple[str, str]], size: int = 700) -> list[tuple[str, str]]:
    chunks: list[tuple[str, str]] = []
    for locator, content in sections:
        for offset, chunk in enumerate(chunk_text(content, size), start=1):
            chunks.append((f"{locator} / 切片 {offset}", chunk))
    return chunks or [("文件级引用", "未识别到可解析正文，需人工补充资料摘要。")]


def extract_sections(storage_path: str | None, mime_type: str | None) -> list[tuple[str, str]]:
    if not storage_path:
        return [("人工登记", "人工登记资料，无上传文件正文。")]
    path = storage.download_to_temp(storage_path)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx_sections(path)
    if suffix in {".xlsx", ".xlsm"}:
        return extract_xlsx_sections(path)
    if suffix in {".txt", ".md", ".csv"}:
        return [(path.name, path.read_text(encoding="utf-8", errors="ignore"))]
    if suffix == ".pdf" or (mime_type and "pdf" in mime_type):
        return extract_pdf_sections(path)
    return [(path.name, path.read_text(encoding="utf-8", errors="ignore"))]


def extract_docx_sections(path: Path) -> list[tuple[str, str]]:
    try:
        doc = Document(path)
        sections = [(f"段落 {index}", paragraph.text) for index, paragraph in enumerate(doc.paragraphs, start=1) if paragraph.text.strip()]
        table_rows: list[tuple[str, str]] = []
        for table_index, table in enumerate(doc.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    table_rows.append((f"表格 {table_index} 行 {row_index}", " | ".join(values)))
        return sections + table_rows
    except Exception:
        with zipfile.ZipFile(path) as archive:
            xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        return [(path.name, re.sub(r"<[^>]+>", " ", xml))]


def extract_xlsx_sections(path: Path) -> list[tuple[str, str]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    sections: list[tuple[str, str]] = []
    for sheet in workbook.worksheets[:5]:
        for row in sheet.iter_rows(max_row=80, values_only=True):
            values = [str(value) for value in row if value is not None]
            if values:
                sections.append((f"工作表 {sheet.title}", " | ".join(values)))
    return sections


def extract_pdf_sections(path: Path) -> list[tuple[str, str]]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        sections = []
        for index, page in enumerate(reader.pages, start=1):
            text_value = page.extract_text() or ""
            if text_value.strip():
                sections.append((f"第 {index} 页", text_value))
        return sections or [(path.name, "PDF 已接收但未识别到文字层，需进入 OCR/VLM 解析。")]
    except Exception:
        return [(path.name, "PDF 已接收。当前运行环境缺少 PDF 文字抽取能力，保留文件级引用。")]


def parse_document(db: Session, document_id: str, job_id: str | None = None) -> dict | None:
    document = db.get(ProjectDocument, document_id)
    if not document:
        return None

    job_id = job_id or f"JOB-{int(time.time() * 1000)}"
    chunks = chunk_sections(extract_sections(document.storage_path, document.mime_type))

    job = db.get(ParseJob, job_id)
    if job:
        job.status = "completed"
        job.message = "资料解析已完成，抽取结果待人工确认。"
        job.result = {"extractedChunks": len(chunks)}
    else:
        db.add(ParseJob(id=job_id, document_id=document_id, status="completed", message="资料解析已完成，抽取结果待人工确认。", result={"extractedChunks": len(chunks)}))
    db.execute(delete(DocumentChunk).where(DocumentChunk.document_id == document_id))
    for index, (locator, content) in enumerate(chunks, start=1):
        db.add(
            DocumentChunk(
                id=f"CHK-{document_id}-{index}",
                project_id=document.project_id,
                document_id=document_id,
                chunk_index=index,
                content=content,
                locator=locator,
            )
        )
    document.parse_status = "已解析"
    document.updated_at = db.scalar(select(text("to_char(current_date, 'YYYY-MM-DD')")))
    db.commit()
    return {"job": {"id": job_id, "status": "completed"}, "document": document, "chunks": len(chunks)}
